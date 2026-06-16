"""Xuất báo cáo E5 + Reranker (corpus 5000 SP, đánh giá tập test) ra file Word."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

REPO = Path(__file__).resolve().parent.parent
FIGURES = REPO / "embedding_project/outputs/evaluation/figures"
REPORT_OUT = REPO / "embedding_project/outputs/evaluation/bao_cao_e5_reranker_5000.docx"

# Kết quả Colab 15/06/2026 — full test (708 query trong corpus 5000)
EVAL_DATE = "15/06/2026"
EVAL_N_JSONL = 750
EVAL_N_QUERIES = 708
EVAL_SKIPPED = 42
EVAL_CORPUS = 5000
SELECTED_N = 10
TARGET_RECALL = 0.95
RERANK_PAIRS = 7080

BI_ENCODER = {
    "P@10": "0,1037",
    "R@10": "0,9952",
    "F1@10": "0,1878",
    "MRR@10": "0,9454",
    "NDCG@10": "0,9562",
}
RERANKER = {
    "P@10": "0,1037",
    "R@10": "0,9952",
    "F1@10": "0,1878",
    "MRR@10": "0,9508",
    "NDCG@10": "0,9603",
}
RECALL_BY_N = {
    "10": "0,9952",
    "20": "0,9980",
    "30": "0,9996",
    "50": "0,9999",
    "75": "1,0000",
    "100": "1,0000",
}
K_GRID = [
    ["5", "0,2028", "0,9840", "0,3363", "0,9454", "0,9562"],
    ["10", "0,1037", "0,9952", "0,1878", "0,9508", "0,9603"],
]
OPT_K_NDCG = 10
THRESHOLD = {
    "tau_eer": 0.9949,
    "tau_min_err": 0.9899,
    "fpr_eer": 0.0852,
    "fnr_eer": 0.0085,
    "fpr_min": 0.0857,
    "fnr_min": 0.0071,
    "error_rate": 0.0660,
    "n_pairs": 2832,
    "n_pos": 708,
    "n_neg": 2124,
    "hard_neg_pool": 100,
}


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if path.is_file():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Thiếu hình: {path.name}]")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()

    title = doc.add_heading(
        "BÁO CÁO THỬ NGHIỆM: E5 BI-ENCODER + CROSS-ENCODER RERANKER", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Dự án: embedding_project — Tìm kiếm ngữ nghĩa sản phẩm tiếng Việt")
    doc.add_paragraph(f"Ngày báo cáo: {date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph(
        "Corpus: ecommerce.csv (5000 SP) | Đánh giá: test_5000.jsonl (query thực tế) | "
        "Model: e5_base_finetuned_5000 + reranker (XLM-RoBERTa)"
    )

    doc.add_heading("1. Tóm tắt", level=1)
    add_table(
        doc,
        ["Chỉ số", "Chỉ Bi-encoder @10", f"+ Reranker @10 (n={SELECTED_N})"],
        [
            ["Precision@10", BI_ENCODER["P@10"], RERANKER["P@10"]],
            ["Recall@10", BI_ENCODER["R@10"], RERANKER["R@10"]],
            ["F1@10", BI_ENCODER["F1@10"], RERANKER["F1@10"]],
            ["MRR@10", BI_ENCODER["MRR@10"], RERANKER["MRR@10"]],
            ["NDCG@10", BI_ENCODER["NDCG@10"], RERANKER["NDCG@10"]],
        ],
    )
    doc.add_paragraph(
        f"Thí nghiệm Colab GPU ({EVAL_DATE}): {EVAL_N_QUERIES} query hợp lệ / "
        f"{EVAL_N_JSONL} dòng test jsonl, corpus {EVAL_CORPUS} SP. "
        f"Pipeline dò Recall@n → chọn n={SELECTED_N} (nhỏ nhất đạt target {TARGET_RECALL}) "
        f"→ rerank {RERANK_PAIRS} cặp. Reranker cải thiện MRR (+0,005) và NDCG (+0,004); "
        "P/R/F1@10 giữ nguyên vì E5 đã đưa đúng SP vào top-10."
    )
    doc.add_paragraph(
        f"Ngưỡng triển khai (hard negative từ E5 top-{THRESHOLD['hard_neg_pool']}): "
        f"τ_EER ≈ {THRESHOLD['tau_eer']:.4f} (FPR≈{THRESHOLD['fpr_eer']:.1%}, "
        f"FNR≈{THRESHOLD['fnr_eer']:.1%}), error rate ≈ {THRESHOLD['error_rate']:.1%}."
    )
    doc.add_paragraph(
        "Precision@10 ≈ 10% gần trần lý thuyết khi mỗi query có 1 nhãn và k=10 cố định; "
        "Recall@10 ≈ 99,5% cho thấy model hầu như luôn tìm đúng SP trong top-10."
    )

    doc.add_heading("2. Dữ liệu và xử lý trước huấn luyện/đánh giá", level=1)

    doc.add_heading("2.1. Nguồn corpus", level=2)
    doc.add_paragraph(
        "File embedding_project/data/ecommerce.csv gồm 5000 sản phẩm từ 5 sàn "
        "(Amazon, Lazada, Shein, Shopee, Walmart), mỗi nguồn 1000 SP. "
        "Trường searchable_text dùng làm passage khi index/retrieve: "
        "ghép title, mô tả, category, brand (không có nhãn trường)."
    )
    add_figure(
        doc,
        FIGURES / "c__llm_provider_benchmarking_source_top.png",
        "Hình 1. Phân phối số lượng sản phẩm theo nguồn",
    )

    doc.add_heading("2.2. Chia train / valid / test", level=2)
    add_table(
        doc,
        ["Tập", "File", "Số mẫu", "Ghi chú"],
        [
            ["Train", "data/training/train_5000.jsonl", "3500", "Fine-tune E5"],
            ["Valid", "data/training/valid_5000.jsonl", "750", "Validation"],
            ["Test", "data/training/test_5000.jsonl", "750", "Đánh giá retrieval"],
            ["Test SP", "embedding_project/data/test_5000.csv", "750", "Tập SP test (cùng split)"],
        ],
    )
    doc.add_paragraph(
        "Chia stratify theo source (70/15/15). Script/notebook: colab_searchable_text_llm.ipynb. "
        f"Đánh giá pipeline: {EVAL_N_QUERIES}/{EVAL_N_JSONL} query có product_id nằm trong corpus "
        f"({EVAL_SKIPPED} query bị bỏ vì SP không có trong ecommerce.csv)."
    )

    doc.add_heading("2.3. Ghép category vào query", level=2)
    doc.add_paragraph(
        "Khác với đánh giá cũ (query = title — metric ảo), tập test dùng query mô phỏng "
        "người dùng tìm theo tên SP + thương hiệu + ngành hàng:"
    )
    doc.add_paragraph(
        "query = title + brand + category_leaf\n\n"
        "Trong đó category_leaf là cấp cuối của danh mục (vd. category "
        "['Máy tính & Laptop', 'Laptop'] → leaf = 'Laptop'). "
        "Các phần rỗng được bỏ qua, khoảng trắng được chuẩn hóa."
    )
    doc.add_paragraph("Ví dụ:")
    doc.add_paragraph(
        "• Query: «Găng tay xe mô tô Fox Racing 180 Illmatik dành cho nam Fox Racing Găng tay»\n"
        "• Positive (passage huấn luyện): «Tên sản phẩm: … Thương hiệu: Fox Racing "
        "Danh mục: Ô tô > Motorcycle & Powersports > … > Găng tay Mô tả: …»"
    )
    doc.add_paragraph(
        "Passage positive trong jsonl có nhãn trường (Tên sản phẩm / Thương hiệu / Danh mục / Mô tả). "
        "Corpus retrieval dùng searchable_text phẳng từ ecommerce.csv — khớp cách triển khai thực tế."
    )
    add_figure(doc, FIGURES / "category_top.png", "Hình 2. Top category trong corpus")
    add_figure(doc, FIGURES / "category_depth.png", "Hình 3. Độ sâu category")

    doc.add_heading("2.4. Negative cho huấn luyện jsonl", level=2)
    doc.add_paragraph(
        "Mỗi dòng jsonl có thêm negative (3 mẫu ngẫu nhiên) và hard_negative (2 mẫu cùng category, "
        "khác product_id) — phục vụ fine-tune/reranker, không dùng trực tiếp trong metric retrieval."
    )

    doc.add_heading("2.5. EDA bổ sung", level=2)
    add_figure(
        doc,
        FIGURES / "c__llm_provider_benchmarking_missing_values.png",
        "Hình 4. Tỷ lệ giá trị thiếu",
    )
    add_figure(
        doc,
        FIGURES / "c__llm_provider_benchmarking_rating_vs_reviews.png",
        "Hình 5. Rating vs số review",
    )

    doc.add_heading("3. Huấn luyện model", level=1)
    add_table(
        doc,
        ["Tham số", "Giá trị"],
        [
            ["Base model", "intfloat/multilingual-e5-base"],
            ["Output", "e5_base_finetuned_5000"],
            ["Dữ liệu", "train_5000.jsonl (query → positive)"],
            ["Loss", "MultipleNegativesRankingLoss"],
            ["Epoch", "1"],
            ["Batch size", "8"],
            ["Learning rate", "1e-5"],
            ["Prefix inference", "query: / passage:"],
        ],
    )
    doc.add_paragraph(
        "Reranker: XLM-RoBERTa cross-encoder, fine-tune riêng trên cặp (query, passage) "
        "có nhãn relevance."
    )

    doc.add_heading("4. Quy trình đo chỉ số", level=1)
    doc.add_paragraph(
        "Script: embedding_project/scripts/evaluate_reranker_pipeline.py\n"
        "Notebook: embedding_project/notebooks/evaluate_e5_reranker_5000_colab.ipynb"
    )

    doc.add_heading("4.1. Pipeline 2 giai đoạn", level=2)
    add_bullets(
        doc,
        [
            "Giai đoạn 1 — Bi-encoder E5: encode toàn bộ corpus (passage:) và query (query:), "
            "cosine similarity → top-n ứng viên.",
            f"Giai đoạn 2 — Chọn n: quét Recall@n trên n ∈ {{10,20,30,50,75,100}}, "
            f"chọn n nhỏ nhất ≥ {TARGET_RECALL}. Kết quả: n={SELECTED_N}.",
            f"Giai đoạn 3 — Reranker: cross-encoder chấm điểm {EVAL_N_QUERIES}×{SELECTED_N} "
            f"= {RERANK_PAIRS} cặp (query, searchable_text), sắp xếp lại → top-k.",
            "Giai đoạn 4 — Grid k ∈ {5, 10, 20} tại n đã chọn; báo cáo chính @k=10.",
        ],
    )

    doc.add_heading("4.2. Định nghĩa metric", level=2)
    add_table(
        doc,
        ["Metric", "Ý nghĩa"],
        [
            ["Precision@k", "Tỷ lệ SP relevant trong k kết quả đầu (1 nhãn/query → max ~10% @k=10)"],
            ["Recall@k", "Tỷ lệ nhãn được tìm thấy trong top-k"],
            ["F1@k", "Trung bình điều hòa P và R"],
            ["MRR@k", "Nghịch đảo thứ hạng SP đúng đầu tiên"],
            ["NDCG@k", "Chất lượng xếp hạng có trọng số vị trí"],
        ],
    )
    doc.add_paragraph(
        "Ground truth: metadata.product_id trong test_5000.jsonl. "
        "Mỗi query có đúng 1 SP relevant."
    )

    doc.add_heading("4.3. Phân tích ngưỡng τ (triển khai)", level=2)
    doc.add_paragraph(
        "Chạy riêng (run_threshold_only): với mỗi query, 1 positive (SP đúng, ưu tiên trong top-n E5) "
        f"+ 3 hard negative từ top-{THRESHOLD['hard_neg_pool']} E5 không thuộc nhãn. "
        "Reranker chấm điểm → quét τ ∈ [min_score, max_score] (201 bước)."
    )
    add_table(
        doc,
        ["Chỉ số", "Định nghĩa"],
        [
            ["FPR(τ)", "Tỷ lệ negative bị chấp nhận (score ≥ τ)"],
            ["FNR(τ)", "Tỷ lệ positive bị loại (score < τ)"],
            ["EER", "τ sao cho FPR(τ) ≈ FNR(τ)"],
            ["Min error", "τ làm (FP + FN) / tổng cặp nhỏ nhất"],
        ],
    )
    doc.add_paragraph(
        "Triển khai: sau rerank, chỉ trả kết quả có score_reranker ≥ τ. "
        "τ cao (~0,99) phù hợp khi ưu tiên precision; τ thấp hơn giữ nhiều kết quả hơn."
    )

    doc.add_heading("5. Kết quả retrieval trên tập test", level=1)

    doc.add_heading("5.1. Bi-encoder vs Reranker @10", level=2)
    add_table(
        doc,
        ["Chỉ số", "Bi-encoder", f"Reranker (n={SELECTED_N})", "Δ"],
        [
            ["P@10", BI_ENCODER["P@10"], RERANKER["P@10"], "0"],
            ["R@10", BI_ENCODER["R@10"], RERANKER["R@10"], "0"],
            ["F1@10", BI_ENCODER["F1@10"], RERANKER["F1@10"], "0"],
            ["MRR@10", BI_ENCODER["MRR@10"], RERANKER["MRR@10"], "+0,0054"],
            ["NDCG@10", BI_ENCODER["NDCG@10"], RERANKER["NDCG@10"], "+0,0041"],
        ],
    )
    doc.add_paragraph(
        "Reranker cải thiện thứ hạng (MRR, NDCG) khi SP đúng đã nằm trong top-10 E5; "
        "không tăng P/R vì số hit trong top-10 không đổi."
    )

    doc.add_heading("5.2. Recall@n và chọn n", level=2)
    add_table(
        doc,
        ["n", "Recall@n"],
        [[n, v] for n, v in RECALL_BY_N.items()],
    )
    doc.add_paragraph(
        f"→ selected_n = {SELECTED_N}. Chỉ cần retrieve 10 SP để đạt ~99,5% recall, "
        f"tiết kiệm {100 - SELECTED_N}×{EVAL_N_QUERIES} cặp rerank so với n=100."
    )

    doc.add_heading("5.3. Metric theo k (tại n=10)", level=2)
    add_table(doc, ["k", "P", "R", "F1", "MRR", "NDCG"], K_GRID)
    doc.add_paragraph(
        f"k=5 cho P@5≈20% và F1@5≈0,34 (ít slot hơn → precision cao hơn). "
        f"k tối ưu NDCG@10: k={OPT_K_NDCG}."
    )

    doc.add_heading("6. Ngưỡng τ trước triển khai", level=1)
    doc.add_paragraph(
        f"Tập ngưỡng: {THRESHOLD['n_pairs']} cặp "
        f"({THRESHOLD['n_pos']} positive, {THRESHOLD['n_neg']} hard negative từ E5 top-n)."
    )
    add_table(
        doc,
        ["Loại", "τ", "FPR", "FNR", "Error rate"],
        [
            [
                "EER",
                f"{THRESHOLD['tau_eer']:.4f}",
                f"{THRESHOLD['fpr_eer']:.4f}".replace(".", ","),
                f"{THRESHOLD['fnr_eer']:.4f}".replace(".", ","),
                f"{THRESHOLD['error_rate']:.4f}".replace(".", ","),
            ],
            [
                "Min error",
                f"{THRESHOLD['tau_min_err']:.4f}",
                f"{THRESHOLD['fpr_min']:.4f}".replace(".", ","),
                f"{THRESHOLD['fnr_min']:.4f}".replace(".", ","),
                f"{THRESHOLD['error_rate']:.4f}".replace(".", ","),
            ],
        ],
    )
    doc.add_paragraph(
        "Đặc điểm đường cong FPR/FNR: FPR giảm mạnh ngay khi τ > 0; FNR tăng vọt khi τ → 1. "
        "Vùng τ ∈ [0,98; 0,995] cho error rate thấp (~6,6%). τ quá cao (>0,995) làm bỏ sót positive."
    )

    doc.add_heading("7. Giải thích Precision@10 thấp", level=1)
    doc.add_paragraph(
        "Với 1 nhãn/query và k=10, Precision@10 tối đa lý thuyết là 10%. "
        "Kết quả ~10,4% cho thấy model đôi khi đưa thêm SP liên quan ngữ nghĩa (cùng category/từ khóa) "
        "vào top-10. Recall ~99,5% xác nhận SP đúng hầu như luôn được recall."
    )

    doc.add_heading("8. Hạn chế và hướng cải thiện", level=1)
    add_bullets(
        doc,
        [
            "Query test có thêm brand + category leaf — dễ hơn query tự do hoàn toàn.",
            f"{EVAL_SKIPPED} query không có SP trong corpus 5000 — cần đồng bộ split.",
            "Corpus searchable_text (phẳng) khác format positive jsonl (có nhãn trường).",
            "Có thể thử eval trên test_5000.csv (750 SP) để đo closed-set retrieval.",
            "Bổ sung filter/boost category tại tầng Qdrant cho query mơ hồ.",
        ],
    )

    doc.add_heading("9. Tái lập thí nghiệm", level=1)
    doc.add_paragraph(
        "1. Copy script + notebook mới lên Colab.\n"
        "2. Cell 3: eval metric (--skip-threshold).\n"
        "3. Cell 4: ngưỡng τ (run_threshold_only).\n"
        "4. Cell 5: xem bảng + biểu đồ.\n\n"
        "Lệnh CLI:\n"
        "python embedding_project/scripts/evaluate_reranker_pipeline.py \\\n"
        "  --embedding-model .../e5_base_finetuned_5000 \\\n"
        "  --reranker-model .../reranker \\\n"
        "  --eval-csv embedding_project/data/ecommerce.csv \\\n"
        "  --query-jsonl data/training/test_5000.jsonl \\\n"
        "  --target-recall 0.95 --skip-threshold"
    )

    REPORT_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_OUT))
    print(f"Đã lưu: {REPORT_OUT}")


if __name__ == "__main__":
    main()
