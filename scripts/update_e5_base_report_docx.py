"""Tạo bao_cao_thu_nghiem_e5_base_v2.docx: pretrained + loss + fine-tune 1/2 epoch."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

METRICS_1EP = Path("embedding_project/outputs/evaluation/metrics_e5_base.json")
METRICS_2EP = Path("embedding_project/outputs/evaluation/metrics_e5_base_2epochs.json")
REPORT_PATH = Path("embedding_project/outputs/evaluation/bao_cao_thu_nghiem_e5_base_v2.docx")

# Loss từ log Colab (MultipleNegativesRankingLoss)
TRAIN_LOSS_1EP = [
    {"epoch": 1, "train_loss": 0.006353, "valid_loss": 0.013850},
]
TRAIN_LOSS_2EP = [
    {"epoch": 1, "train_loss": 0.006079, "valid_loss": 0.013244},
    {"epoch": 2, "train_loss": 0.019181, "valid_loss": 0.010515},
]


def fmt(x: float, digits: int = 3) -> str:
    return f"{x:.{digits}f}"


def fmt_loss(x: float) -> str:
    return f"{x:.6f}"


def pct_delta(before: float, after: float) -> str:
    if before == 0:
        return "—"
    return f"{((after - before) / before) * 100:+.1f}%"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def loss_rows(epochs: list[dict]) -> list[list[str]]:
    return [
        [
            str(e["epoch"]),
            fmt_loss(e["train_loss"]),
            fmt_loss(e["valid_loss"]),
        ]
        for e in epochs
    ]


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def main() -> None:
    data1 = load_json(METRICS_1EP)
    data2 = load_json(METRICS_2EP)
    pre = data1["pretrained"]
    ft1 = data1["finetuned"]
    ft2 = data2["finetuned"]

    doc = Document()

    title = doc.add_heading("BÁO CÁO THỬ NGHIỆM MÔ HÌNH intfloat/multilingual-e5-base", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Dự án: embedding_project")
    doc.add_paragraph("Ngày cập nhật: 04/06/2026")

    doc.add_heading("1. Mục tiêu", level=1)
    doc.add_paragraph(
        "Đánh giá truy hồi ngữ nghĩa E5-base: pretrained, fine-tune 1 epoch và 2 epoch "
        "(loss huấn luyện + metric retrieval @10)."
    )

    doc.add_heading("2. Cấu hình thử nghiệm", level=1)
    for item in [
        "Model gốc: intfloat/multilingual-e5-base",
        "Fine-tune 1 epoch: embedding_project/models/e5_base_finetuned_final/",
        "Fine-tune 2 epoch: embedding_project/models/e5_base_finetuned_2ep_final/",
        "Loss: MultipleNegativesRankingLoss (contrastive, không phải perplexity LLM)",
        "Quy ước retrieval: query: / passage:",
        "Dữ liệu: train ~6.224, valid ~778, test ~779 cặp query–positive",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Quá trình huấn luyện (Training / Validation Loss)", level=1)

    doc.add_heading("3.1. Fine-tune 1 epoch", level=2)
    add_table(
        doc,
        ["Epoch", "Training Loss", "Validation Loss"],
        loss_rows(TRAIN_LOSS_1EP),
    )

    doc.add_heading("3.2. Fine-tune 2 epoch", level=2)
    add_table(
        doc,
        ["Epoch", "Training Loss", "Validation Loss"],
        loss_rows(TRAIN_LOSS_2EP),
    )

    doc.add_paragraph("Nhận xét loss:")
    e1_2 = TRAIN_LOSS_2EP[0]
    e2_2 = TRAIN_LOSS_2EP[1]
    for item in [
        f"1 epoch: train={TRAIN_LOSS_1EP[0]['train_loss']:.6f}, valid={TRAIN_LOSS_1EP[0]['valid_loss']:.6f}.",
        f"2 epoch — Epoch 1: train={e1_2['train_loss']:.6f}, valid={e1_2['valid_loss']:.6f}.",
        f"2 epoch — Epoch 2: train={e2_2['train_loss']:.6f}, valid={e2_2['valid_loss']:.6f} "
        f"(valid giảm {pct_delta(e1_2['valid_loss'], e2_2['valid_loss'])} so với epoch 1).",
        "Epoch 2: training loss tăng (0.019) nhưng validation loss thấp nhất — "
        "load_best_model_at_end chọn checkpoint epoch 2 theo eval_loss.",
        "Không có dấu hiệu overfit nặng trên valid: valid loss epoch 2 < epoch 1.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Kết quả pretrained (retrieval @10)", level=1)
    add_table(
        doc,
        ["Chỉ số", "Pretrained"],
        [
            ["Precision@10", fmt(pre["Precision@10"], 6)],
            ["Recall@10", fmt(pre["Recall@10"], 6)],
            ["MRR@10", fmt(pre["MRR@10"], 6)],
            ["NDCG@10", fmt(pre["NDCG@10"], 6)],
        ],
    )

    doc.add_heading("5. Fine-tuned — 1 epoch (e5_base_finetuned_final)", level=1)
    doc.add_heading("5.1. Bảng metric retrieval", level=2)
    add_table(
        doc,
        ["Metric", "1 epoch", "Δ (vs pretrained)"],
        [
            ["NDCG@10", fmt(ft1["NDCG@10"]), pct_delta(pre["NDCG@10"], ft1["NDCG@10"])],
            ["Precision@10", fmt(ft1["Precision@10"]), pct_delta(pre["Precision@10"], ft1["Precision@10"])],
            ["Recall@10", fmt(ft1["Recall@10"]), pct_delta(pre["Recall@10"], ft1["Recall@10"])],
            ["MRR@10", fmt(ft1["MRR@10"]), pct_delta(pre["MRR@10"], ft1["MRR@10"])],
        ],
    )
    doc.add_heading("5.2. Nhận xét", level=2)
    for item in [
        f"NDCG@10: {pre['NDCG@10']:.3f} → {ft1['NDCG@10']:.3f}.",
        f"Recall@10: {ft1['Recall@10']:.1%}.",
        f"MRR@10: {ft1['MRR@10']:.3f}.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Fine-tuned — 2 epoch (e5_base_finetuned_2ep_final)", level=1)
    doc.add_heading("6.1. Bảng metric retrieval", level=2)
    add_table(
        doc,
        ["Metric", "2 epoch", "Δ (vs pretrained)"],
        [
            ["NDCG@10", fmt(ft2["NDCG@10"]), pct_delta(pre["NDCG@10"], ft2["NDCG@10"])],
            ["Precision@10", fmt(ft2["Precision@10"]), pct_delta(pre["Precision@10"], ft2["Precision@10"])],
            ["Recall@10", fmt(ft2["Recall@10"]), pct_delta(pre["Recall@10"], ft2["Recall@10"])],
            ["MRR@10", fmt(ft2["MRR@10"]), pct_delta(pre["MRR@10"], ft2["MRR@10"])],
        ],
    )
    doc.add_heading("6.2. Nhận xét", level=2)
    for item in [
        f"NDCG@10: {ft2['NDCG@10']:.3f} — cao hơn 1 epoch ({pct_delta(ft1['NDCG@10'], ft2['NDCG@10'])}).",
        f"Recall@10: {ft2['Recall@10']:.1%} (1 epoch: {ft1['Recall@10']:.1%}).",
        f"MRR@10: {ft2['MRR@10']:.3f} ({pct_delta(ft1['MRR@10'], ft2['MRR@10'])} vs 1 epoch).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. So sánh 1 epoch vs 2 epoch (retrieval)", level=1)
    add_table(
        doc,
        ["Metric", "1 epoch", "2 epoch", "Δ (2 vs 1)"],
        [
            ["NDCG@10", fmt(ft1["NDCG@10"]), fmt(ft2["NDCG@10"]), pct_delta(ft1["NDCG@10"], ft2["NDCG@10"])],
            ["Precision@10", fmt(ft1["Precision@10"]), fmt(ft2["Precision@10"]), pct_delta(ft1["Precision@10"], ft2["Precision@10"])],
            ["Recall@10", fmt(ft1["Recall@10"]), fmt(ft2["Recall@10"]), pct_delta(ft1["Recall@10"], ft2["Recall@10"])],
            ["MRR@10", fmt(ft1["MRR@10"]), fmt(ft2["MRR@10"]), pct_delta(ft1["MRR@10"], ft2["MRR@10"])],
        ],
    )
    add_table(
        doc,
        ["Epoch", "Train loss (1 ep run)", "Valid loss (1 ep run)", "Train loss (2 ep run)", "Valid loss (2 ep run)"],
        [
            [
                "1",
                fmt_loss(TRAIN_LOSS_1EP[0]["train_loss"]),
                fmt_loss(TRAIN_LOSS_1EP[0]["valid_loss"]),
                fmt_loss(TRAIN_LOSS_2EP[0]["train_loss"]),
                fmt_loss(TRAIN_LOSS_2EP[0]["valid_loss"]),
            ],
            [
                "2",
                "—",
                "—",
                fmt_loss(TRAIN_LOSS_2EP[1]["train_loss"]),
                fmt_loss(TRAIN_LOSS_2EP[1]["valid_loss"]),
            ],
        ],
    )

    metrics_keys = ["NDCG@10", "Precision@10", "Recall@10", "MRR@10"]
    better_2 = sum(1 for k in metrics_keys if ft2[k] > ft1[k])
    doc.add_paragraph(
        f"Retrieval: 2 epoch tốt hơn 1 epoch trên {better_2}/4 metric. "
        f"Loss: valid epoch 2 (0.010515) thấp nhất trong run 2 epoch."
    )

    doc.add_heading("8. Kết luận", level=1)
    for item in [
        "E5-base pretrained là baseline mạnh; fine-tune cải thiện retrieval so với pretrained.",
        "2 epoch: valid loss giảm ở epoch 2; metric retrieval (NDCG, Recall) tốt hơn 1 epoch.",
        "Model đề xuất triển khai: e5_base_finetuned_2ep_final (metric cao hơn).",
        "Model 1 epoch: e5_base_finetuned_final/ (nhẹ hơn về thời gian train).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    out = REPORT_PATH
    try:
        doc.save(str(out))
    except PermissionError:
        out = REPORT_PATH.with_name(REPORT_PATH.stem + "_updated.docx")
        doc.save(str(out))
        print(f"File gốc đang mở — đã ghi bản mới: {out}")
        return
    print(f"Đã ghi: {out}")


if __name__ == "__main__":
    main()
