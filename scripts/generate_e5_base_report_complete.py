"""Báo cáo hoàn chỉnh E5-base: Pretrained vs Fine-tune 2 epoch + Learning Curve."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

REPO = Path(__file__).resolve().parent.parent
METRICS_PRE_FT = REPO / "embedding_project/outputs/evaluation/metrics_e5_base.json"
METRICS_2EP = REPO / "embedding_project/outputs/evaluation/metrics_e5_base_2epochs.json"
LEARNING_CURVE = REPO / "embedding_project/outputs/evaluation/e5_base_learning_curve_6epoch.png"
REPORT_OUT = REPO / "embedding_project/outputs/evaluation/bao_cao_thu_nghiem_e5_base_hoan_chinh.docx"

TRAIN_LOSS_2EP = [
    {"epoch": 1, "train_loss": 0.006079, "valid_loss": 0.013244},
    {"epoch": 2, "train_loss": 0.019181, "valid_loss": 0.010515},
]

# Metric fine-tune 2 epoch (Colab eval test_cleaned, làm tròn 4 chữ số)
FT2 = {
    "Precision@10": 0.1332,
    "Recall@10": 0.9857,
    "MRR@10": 0.8940,
    "NDCG@10": 0.9143,
}

# Category pollution@10: % kết quả top-10 thuộc category khác positive
CATEGORY_POLLUTION = {
    "pretrained": 0.810,
    "finetuned_2ep": 0.794,
}


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(x: float, digits: int = 4) -> str:
    return f"{x:.{digits}f}"


def fmt_pct(x: float) -> str:
    return f"{x:.2%}"


def pct_delta(before: float, after: float) -> str:
    if before == 0:
        return "—"
    return f"{((after - before) / before) * 100:+.1f}%"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    data = load_json(METRICS_PRE_FT)
    pre = data["pretrained"]

    doc = Document()

    title = doc.add_heading(
        "BÁO CÁO THỬ NGHIỆM MÔ HÌNH EMBEDDING E5-BASE", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Dự án: embedding_project — Semantic Product Search (tiếng Việt)")
    doc.add_paragraph(f"Ngày báo cáo: {date.today().strftime('%d/%m/%Y')}")

    doc.add_heading("1. Mục tiêu", level=1)
    doc.add_paragraph(
        "Đánh giá chất lượng truy hồi sản phẩm bằng embedding intfloat/multilingual-e5-base "
        "trên bộ dữ liệu query–product tiếng Việt. So sánh model pretrained với model "
        "fine-tune 2 epoch; phân tích quá trình huấn luyện qua loss và learning curve."
    )

    doc.add_heading("2. Cấu hình thử nghiệm", level=1)
    for item in [
        "Model gốc (pretrained): intfloat/multilingual-e5-base",
        "Model fine-tune (đề xuất triển khai): embedding_project/models/e5_base_finetuned_2ep_final/",
        "Loss: MultipleNegativesRankingLoss (contrastive retrieval)",
        "Quy ước E5: prefix query: (truy vấn) và passage: (corpus sản phẩm)",
        "Huấn luyện: 2 epoch, load_best_model_at_end=True (chọn checkpoint valid loss thấp nhất)",
        "Đánh giá: retrieval trên tập test, k=10",
        "Corpus: merged_products_vi_cleaned.csv (~2.000 sản phẩm)",
        "Dữ liệu train / valid / test: ~6.224 / ~778 / ~779 cặp query–positive",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Quá trình huấn luyện (Fine-tune 2 epoch)", level=1)

    doc.add_heading("3.1. Training / Validation Loss", level=2)
    add_table(
        doc,
        ["Epoch", "Training Loss", "Validation Loss"],
        [
            [str(e["epoch"]), f"{e['train_loss']:.6f}", f"{e['valid_loss']:.6f}"]
            for e in TRAIN_LOSS_2EP
        ],
    )

    doc.add_paragraph("Nhận xét loss:")
    e1, e2 = TRAIN_LOSS_2EP
    for item in [
        f"Epoch 1: train={e1['train_loss']:.6f}, valid={e1['valid_loss']:.6f}.",
        f"Epoch 2: train={e2['train_loss']:.6f}, valid={e2['valid_loss']:.6f} "
        f"(valid giảm {pct_delta(e1['valid_loss'], e2['valid_loss'])} so với epoch 1).",
        "Training loss epoch 2 tăng (0.019) trong khi validation loss giảm — "
        "checkpoint epoch 2 được chọn làm model cuối nhờ load_best_model_at_end.",
        "Không có dấu hiệu overfit nặng trên valid tại epoch 2.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2. Learning Curve (thử nghiệm mở rộng 6 epoch)", level=2)
    doc.add_paragraph(
        "Để kiểm tra liệu huấn luyện thêm epoch có lợi hay không, đã chạy thử nghiệm "
        "fine-tune 6 epoch trên cùng dữ liệu. Biểu đồ dưới đây cho thấy valid loss thấp nhất "
        "tại epoch 2 (đường kẻ xanh best valid), xác nhận lựa chọn dừng ở 2 epoch là hợp lý."
    )
    if LEARNING_CURVE.is_file():
        doc.add_picture(str(LEARNING_CURVE), width=Inches(5.8))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Hình 1. Learning curve E5-base (train_loss / valid_loss, 6 epoch)")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Thiếu ảnh: {LEARNING_CURVE}]")

    doc.add_heading("4. Kết quả đánh giá retrieval (@10)", level=1)

    doc.add_heading("4.1. Pretrained (intfloat/multilingual-e5-base)", level=2)
    add_table(
        doc,
        ["Chỉ số", "Giá trị"],
        [
            ["Precision@10", fmt(pre["Precision@10"])],
            ["Recall@10", fmt(pre["Recall@10"])],
            ["MRR@10", fmt(pre["MRR@10"])],
            ["NDCG@10", fmt(pre["NDCG@10"])],
        ],
    )

    doc.add_heading("4.2. Fine-tuned 2 epoch 🏆 (e5_base_finetuned_2ep_final)", level=2)
    add_table(
        doc,
        ["Chỉ số", "Giá trị"],
        [
            ["Precision@10", fmt(FT2["Precision@10"])],
            ["Recall@10", fmt(FT2["Recall@10"])],
            ["MRR@10", fmt(FT2["MRR@10"])],
            ["NDCG@10", fmt(FT2["NDCG@10"])],
        ],
    )

    doc.add_heading("4.3. So sánh Pretrained vs Fine-tune 2 epoch", level=2)
    pollution_delta_pp = (
        (CATEGORY_POLLUTION["finetuned_2ep"] - CATEGORY_POLLUTION["pretrained"]) * 100
    )
    add_table(
        doc,
        ["Metric", "Pretrained", "Fine-tune 2 epoch", "Δ (2ep vs pre)"],
        [
            [
                "Precision@10",
                fmt(pre["Precision@10"]),
                fmt(FT2["Precision@10"]),
                pct_delta(pre["Precision@10"], FT2["Precision@10"]),
            ],
            [
                "Recall@10",
                fmt(pre["Recall@10"]),
                fmt(FT2["Recall@10"]),
                pct_delta(pre["Recall@10"], FT2["Recall@10"]),
            ],
            [
                "MRR@10",
                fmt(pre["MRR@10"]),
                fmt(FT2["MRR@10"]),
                pct_delta(pre["MRR@10"], FT2["MRR@10"]),
            ],
            [
                "NDCG@10",
                fmt(pre["NDCG@10"]),
                fmt(FT2["NDCG@10"]),
                pct_delta(pre["NDCG@10"], FT2["NDCG@10"]),
            ],
            [
                "Avg Category pollution@10",
                fmt_pct(CATEGORY_POLLUTION["pretrained"]),
                fmt_pct(CATEGORY_POLLUTION["finetuned_2ep"]),
                f"{pollution_delta_pp:+.1f} điểm (âm = 2ep ít lẫn category hơn)",
            ],
        ],
    )

    doc.add_paragraph("Tóm tắt cải thiện sau fine-tune 2 epoch:")
    for item in [
        f"NDCG@10: {pre['NDCG@10']:.4f} → {FT2['NDCG@10']:.4f} ({pct_delta(pre['NDCG@10'], FT2['NDCG@10'])}).",
        f"Recall@10: {fmt_pct(pre['Recall@10'])} → {fmt_pct(FT2['Recall@10'])} — "
        f"gần như tìm được sản phẩm liên quan trong top-10.",
        f"MRR@10: {pre['MRR@10']:.4f} → {FT2['MRR@10']:.4f} ({pct_delta(pre['MRR@10'], FT2['MRR@10'])}).",
        f"Precision@10: {pct_delta(pre['Precision@10'], FT2['Precision@10'])} — "
        "tỷ lệ hit trong top-10 tăng nhẹ nhưng vẫn thấp (~13%).",
        f"Category pollution: {fmt_pct(CATEGORY_POLLUTION['pretrained'])} → "
        f"{fmt_pct(CATEGORY_POLLUTION['finetuned_2ep'])} "
        f"({pollution_delta_pp:+.1f} điểm phần trăm) — cải thiện nhẹ, chưa đủ cho production.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.4. Đánh giá thủ công trên query thực tế (OOD)", level=2)
    doc.add_paragraph(
        "Metric trên test_cleaned phản ánh phân phối train (query sinh từ attribute/category). "
        "Đánh giá thủ công với query người dùng thực tế cho thấy nhiều trường hợp fine-tune "
        "không tốt hơn — thậm chí tệ hơn — pretrained."
    )
    add_table(
        doc,
        ["Query thực tế", "Pretrained (top)", "Fine-tune 2ep (top)", "Vấn đề"],
        [
            [
                "tai nghe chống ồn",
                "#2: tai nghe khử tiếng ồn cho chó",
                "#1: tai nghe khử tiếng ồn cho chó (B0CBQW9MF1)",
                "Lexical pollution: title SP chó chứa cụm 'tai nghe khử tiếng ồn'; "
                "category Đồ dùng cho thú cưng, không phải Điện tử > Tai nghe.",
            ],
            [
                "giày chạy bộ nam nhẹ",
                "Ultraboost, Saucony, ASICS (đúng intent)",
                "Sneaker thời trang, giày công nghiệp",
                "Fine-tune học association sai từ label train; "
                "query 'giày chạy bộ' gắn positive giày công nghiệp/sneaker.",
            ],
            [
                "áo mưa nam chống nước",
                "Giày mưa, ô, áo nén",
                "Tương tự pretrained",
                "Corpus không có áo mưa nam — semantic search không thể trả đúng SP không tồn tại.",
            ],
            [
                "giày chạy bộ thoáng khí",
                "MRR = 0",
                "MRR = 0",
                "Label train gán positive sai (giày công nghiệp Wolverine, sneaker SeaVees) — "
                "cả hai model đều fail.",
            ],
        ],
    )
    doc.add_paragraph(
        "Kết quả Qdrant (collection products_vi_e5_2ep) với query 'tai nghe chống ồn': "
        "SP chó score 0.548 (#1), tai nghe Bluetooth chống ồn thật (B0C6JMP9LH) chỉ #5 (0.460)."
    )

    doc.add_heading("5. Điểm yếu và hạn chế", level=1)

    doc.add_heading("5.1. Metric tổng lạc quan so với trải nghiệm thực tế", level=2)
    for item in [
        "Recall@10 ≈ 98.6% và NDCG@10 ≈ 0.914 trên test_cleaned — cao vì test cùng cách sinh query "
        "với train (attribute + category template), không đại diện cho query tự do của người dùng.",
        "Precision@10 chỉ ~13%: trung bình mỗi query có ~1.3 hit đúng trong 10 kết quả — "
        "nhiều slot top-10 bị chiếm bởi SP không liên quan.",
        "Category pollution@10 vẫn ~79%: gần 8/10 kết quả top-10 thuộc category khác positive; "
        "fine-tune chỉ giảm 1.6 điểm phần trăm so với pretrained.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5.2. Chất lượng dữ liệu huấn luyện", level=2)
    for item in [
        "Một số cặp query–positive trong train_cleaned.jsonl gán sai intent: "
        "ví dụ 'giày chạy bộ thoáng khí' → giày công nghiệp (Wolverine B00U4ADYCM), "
        "sneaker SeaVees, giày công trường New Balance 412.",
        "Query sinh tự động theo attribute (màu, size, chất liệu) dễ tạo query mơ hồ, "
        "model học match theo từ khóa surface thay vì intent mua hàng.",
        "Không có hard negative theo category trong loss — "
        "MultipleNegativesRankingLoss chỉ đẩy positive lên, không đẩy SP cùng từ khóa khác category xuống.",
        "Train/valid/test chia random 80/10/10, không stratify theo category — "
        "metric tổng có thể che lỗi trên ngành hàng cụ thể.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5.3. Corpus và mô tả sản phẩm", level=2)
    for item in [
        "Keyword pollution trong title/description: SP bảo vệ tai cho chó dùng từ 'tai nghe khử tiếng ồn' "
        "→ match lexical với query 'tai nghe chống ồn' dù category hoàn toàn khác.",
        "Corpus ~2.000 SP không phủ đủ ngành hàng — query 'áo mưa nam' không có SP tương ứng, "
        "model trả kết quả gần nghĩa nhất (giày mưa, ô).",
        "Một số SP SHEIN bị phân loại sai category (tai nghe TWS nằm trong Đồ Chơi > Thiết Bị Ghi Âm) "
        "→ làm tăng pollution và khó filter theo category.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5.4. Hành vi model sau fine-tune", level=2)
    for item in [
        "Fine-tune khuếch đại association đã học từ train — có thể làm tệ hơn pretrained "
        "trên query OOD (ví dụ 'giày chạy bộ nam nhẹ': pretrained đúng, 2ep sai).",
        "Model không có cơ chế phân biệt category khi query không chỉ rõ (không có 'chó' trong "
        "'tai nghe chống ồn' nhưng vẫn trả SP thú cưng).",
        "Chỉ dựa embedding semantic thuần — không có rerank lexical (BM25) hay filter category ở tầng retrieval.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Hướng cải thiện", level=1)

    doc.add_heading("6.1. Dữ liệu (ưu tiên cao)", level=2)
    add_table(
        doc,
        ["Hạng mục", "Mô tả", "Kỳ vọng"],
        [
            [
                "Audit & làm sạch train",
                "Script clean_train_intent_v2.py — đã loại 153 cặp train (xem Mục 7.1).",
                "Giảm học association sai; metric thủ công cải thiện.",
            ],
            [
                "Hard negative theo category",
                "Mine SP khác category + khác intent, train TripletLoss "
                "(đã triển khai v2 — xem Mục 7).",
                "Giảm category pollution; đẩy SP đúng ngành lên top.",
            ],
            [
                "Normalize text khi index",
                "Với SP pet/industrial: thay 'tai nghe' → 'bảo vệ tai' trong field embed; "
                "giữ title gốc cho hiển thị.",
                "Giảm lexical pollution mà không cần train lại.",
            ],
            [
                "Bộ eval thủ công",
                "~20 query production-critical (tai nghe, giày chạy, áo mưa, …) — "
                "bắt buộc pass trước khi đổi model Qdrant.",
                "Phát hiện regression sớm, bổ sung cho metric test_cleaned.",
            ],
        ],
    )

    doc.add_heading("6.2. Retrieval & triển khai", level=2)
    for item in [
        "Filter/boost category tại Qdrant: query chứa 'tai nghe' → ưu tiên Điện tử > Tai nghe, "
        "loại Đồ dùng cho thú cưng.",
        "Hybrid search: BM25 (từ khóa) + semantic embedding, kết hợp RRF hoặc weighted score.",
        "Rerank top-50 bằng cross-encoder hoặc LLM nhẹ cho query quan trọng.",
        "Metric theo category: báo cáo NDCG/MRR/pollution riêng từng ngành hàng, không chỉ trung bình.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.3. Huấn luyện", level=2)
    for item in [
        "Giữ task query–product + MultipleNegativesRankingLoss (đã thống nhất kiến trúc).",
        "Dừng ở 2 epoch (learning curve 6 epoch xác nhận valid loss tốt nhất tại epoch 2).",
        "Không train thêm epoch trên dữ liệu chưa làm sạch — sẽ khuếch đại label noise.",
        "Sau khi làm sạch data + hard negatives: train lại 2 epoch, so sánh metric tổng và eval thủ công.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Cải tiến pipeline đã triển khai (giai đoạn 2)", level=1)
    doc.add_paragraph(
        "Sau khi phát hiện metric test_cleaned lạc quan nhưng query thực tế vẫn fail, "
        "đã triển khai chuỗi cải tiến: làm sạch train → mine hard negative có kiểm soát → "
        "fine-tune bổ sung bằng TripletLoss từ checkpoint 1 epoch (ít overfit hơn 2 epoch)."
    )

    doc.add_heading("7.1. Làm sạch dữ liệu huấn luyện (train_cleaned_v2)", level=2)
    doc.add_paragraph(
        "Script: embedding_project/scripts/clean_train_intent_v2.py"
    )
    add_table(
        doc,
        ["Tập", "Trước", "Sau", "Đã loại"],
        [
            ["train_cleaned_v2.jsonl", "6.224", "6.071", "153 cặp"],
            ["valid_cleaned_v2.jsonl", "778", "761", "17 cặp"],
        ],
    )
    doc.add_paragraph("Các nhóm cặp bị loại (train):")
    for item in [
        "pet_product_human_query (119): query người nhưng positive là đồ thú cưng.",
        "running_shoe_wrong_category (23): query giày chạy bộ gán positive giày công nghiệp/sneaker.",
        "generic_shoe_work_boot (5), headphone_query_truck_driver (4): pollution điển hình.",
        "Log chi tiết: embedding_project/data/train_removed_pairs_v2.jsonl",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7.2. Mine hard negative theo category (v2)", level=2)
    doc.add_paragraph(
        "Script: embedding_project/scripts/mine_category_hard_negatives_from_qdrant.py. "
        "Mine từ Qdrant bằng vector search, xuất triplets JSONL cho TripletLoss."
    )
    doc.add_paragraph("Pipeline mine v2 — điều kiện negative (tất cả phải thỏa):")
    for item in [
        "product_id ≠ positive; category leaf khác sau normalize (loại sneaker↔sneaker chỉ khác nam/nữ).",
        "Không cùng intent cụ thể với query (tai nghe, chạy bộ, sneaker, máy sấy tóc, …).",
        "Không cùng subtype giày sai (chạy bộ vs sneaker thể thao); giữ giày công nghiệp/mũi thép.",
        "Loại cặp nhầm: tai nghe ↔ bông tai; giày trượt ↔ sneaker.",
        "Score cosine trong band [0.48, 0.84] — không quá dễ, không quá giống (false negative).",
        "Bỏ top-1 search (min_rank=2); mỗi negative_product_id tái sử dụng tối đa 5 lần.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Kết quả mine full (2.000 anchor, neg_per_anchor=1): "
        "180 triplets / 1.820 anchor không tìm được negative hợp lệ. "
        "File: embedding_project/data/train_triplets_category_hardneg_v2.jsonl. "
        "Ưu tiên chất lượng hơn số lượng — tỷ lệ thu được ~9% anchor."
    )

    doc.add_heading("7.3. Huấn luyện TripletLoss (từ model 1 epoch)", level=2)
    doc.add_paragraph(
        "Script: embedding_project/scripts/train_e5_base_triplet_loss_from_jsonl.py. "
        "Notebook Colab: embedding_project/notebooks/train_e5_triplet_hardneg_colab.ipynb."
    )
    add_table(
        doc,
        ["Tham số", "Giá trị"],
        [
            ["Base model", "e5_base_finetuned_final (1 epoch, không train lại từ pretrained)"],
            ["Loss", "TripletLoss, triplet_margin=0.2"],
            ["Triplets", "180 (train_triplets_category_hardneg_v2.jsonl)"],
            ["Epochs / batch / lr", "1 / 16 / 1e-5"],
            ["Output", "e5_base_triplet_hardneg_20260608_0827/"],
        ],
    )
    doc.add_paragraph(
        "Lý do train từ 1 epoch: đã học domain tiếng Việt nhưng ít overfit hơn 2 epoch; "
        "TripletLoss bổ sung tín hiệu đẩy hard negative xuống mà không phá representation cơ bản."
    )
    doc.add_paragraph(
        "Smoke test sau train (1 triplet): query 'thảm sàn an toàn' — "
        "cos(query, positive)=0.73, cos(query, negative)=0.42 → positive gần hơn negative."
    )

    doc.add_heading("7.4. Index Qdrant collection mới", level=2)
    for item in [
        "Collection: products_vi_e5_triplet (2.000 vector, dim 768).",
        "Model encode: e5_base_triplet_hardneg_20260608_0827, prefix query:/passage:.",
        "Lệnh: python vector_db/03_index_to_qdrant.py --recreate "
        "--collection products_vi_e5_triplet --model-path ... --e5-prefix",
        "Cùng server Qdrant với products_vi_e5_2ep; phục vụ A/B search thực tế.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7.5. Đánh giá sơ bộ model Triplet (query dễ nhầm category)", level=2)
    doc.add_paragraph(
        "Test top-5 trên corpus 2.000 SP (không qua Qdrant), 10 query hay lỗi. "
        "File: embedding_project/outputs/evaluation/triplet_confusing_queries_top5.csv"
    )
    add_table(
        doc,
        ["Query", "Top-1 Triplet", "Đánh giá"],
        [
            ["tai nghe không dây trắng", "Disney J300 tai nghe không dây", "Đúng loại"],
            ["giày sneaker thời trang nhẹ", "New Balance 574 sneaker nam", "Đúng category"],
            ["dầu dưỡng tóc giảm rụng", "Dầu dưỡng tóc Adivasi", "Đúng loại"],
            ["ốp lưng điện thoại chống sốc", "Ốp Samsung Galaxy", "Đúng loại"],
            ["giày chạy bộ nam nhẹ", "Sperry sneaker thời trang", "Sai intent (không phải chạy bộ)"],
            ["giày nam", "Wolverine giày ủng công nghiệp", "Pollution còn (giày công nghiệp)"],
            ["tai nghe bluetooth chống ồn", "Tai nghe tài xế xe tải", "Pollution còn"],
            ["máy sấy tóc", "Đèn sấy móng UV", "Sai hoàn toàn"],
            ["son môi màu hồng đất", "Sơn móng tay", "Nhầm son môi ↔ nail"],
        ],
    )
    doc.add_paragraph("Nhận xét giai đoạn 2:")
    for item in [
        "Triplet cải thiện một số case rõ (tai nghe không dây, ốp lưng, dầu tóc) so với hành vi 2ep trên query tương tự.",
        "Các lỗi pollution cứng (giày công nghiệp, máy sấy móng, tai nghe xe tải) vẫn còn — "
        "180 triplets chưa phủ đủ ngành hàng và chưa thay thế normalize corpus.",
        "Top-2/3 đôi khi đúng hơn top-1 (vd. earbud chống ồn ở rank 2–3 cho 'tai nghe bluetooth chống ồn').",
        "Chưa có metric đầy đủ trên test_cleaned cho model triplet — cần eval formal + 78 query manual.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7.6. Hướng tiếp theo sau giai đoạn 2", level=2)
    for item in [
        "Mở rộng triplets: nới score band hoặc mine thêm anchor; có thể dùng NVIDIA LLM filter false negative.",
        "Eval A/B: 1ep vs 2ep vs triplet trên manual_eval_queries.csv (78 query).",
        "Normalize searchable_text khi index (pet/industrial keyword) — không chỉ dựa train.",
        "Hybrid retrieval + filter category tại Qdrant cho query mơ hồ.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. Kết luận và khuyến nghị", level=1)
    for item in [
        "Fine-tune 2 epoch cải thiện rõ rệt metric retrieval trên test_cleaned "
        f"(NDCG +{pct_delta(pre['NDCG@10'], FT2['NDCG@10'])}, "
        f"MRR +{pct_delta(pre['MRR@10'], FT2['MRR@10'])}, "
        f"Recall +{pct_delta(pre['Recall@10'], FT2['Recall@10'])}).",
        "Tuy nhiên, metric tổng KHÔNG đủ điều kiện kết luận 'sẵn sàng production': "
        "category pollution ~79%, precision thấp, nhiều query thực tế fail.",
        "Model 2ep đã index Qdrant (products_vi_e5_2ep) phục vụ thử nghiệm/demo; "
        "chưa nên thay thế hoàn toàn pretrained cho traffic thật cho đến khi pass bộ eval thủ công.",
        "Giai đoạn 2 đã triển khai: làm sạch train v2, mine hard negative v2 (180 triplets), "
        "TripletLoss từ 1ep, index Qdrant products_vi_e5_triplet — xem Mục 7.",
        "Ưu tiên tiếp theo: (1) eval formal triplet vs 1ep/2ep, (2) mở rộng triplets + LLM filter, "
        "(3) normalize corpus khi index, (4) hybrid/category filter tại retrieval.",
        "Learning curve 6 epoch xác nhận epoch 2 là checkpoint tốt nhất; không cần train thêm epoch "
        "trên dữ liệu hiện tại.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Tài liệu tham chiếu", level=1)
    for item in [
        f"Metric pretrained + 1ep: {METRICS_PRE_FT.relative_to(REPO)}",
        f"Metric 2ep: {METRICS_2EP.relative_to(REPO)}",
        "Notebook train 2ep: embedding_project/notebooks/train_embedding_model_e5_base_2epochs.ipynb",
        "Notebook train triplet: embedding_project/notebooks/train_e5_triplet_hardneg_colab.ipynb",
        "Mine hard neg v2: embedding_project/scripts/mine_category_hard_negatives_from_qdrant.py",
        "Train triplet: embedding_project/scripts/train_e5_base_triplet_loss_from_jsonl.py",
        "Triplets v2: embedding_project/data/train_triplets_category_hardneg_v2.jsonl",
        "Model triplet: embedding_project/models/e5_base_triplet_hardneg_20260608_0827/",
        "Index Qdrant 2ep: vector_db/03_index_to_qdrant.py --preset-e5-2ep",
        "Index Qdrant triplet: collection products_vi_e5_triplet",
        "Eval query khó: embedding_project/outputs/evaluation/triplet_confusing_queries_top5.csv",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    REPORT_OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(REPORT_OUT))
        print(f"Đã ghi: {REPORT_OUT}")
    except PermissionError:
        alt = REPORT_OUT.with_name(REPORT_OUT.stem + "_new.docx")
        doc.save(str(alt))
        print(f"File đang mở — đã ghi: {alt}")


if __name__ == "__main__":
    main()
