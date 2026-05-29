"""Generate DOCX evaluation report for first fine-tuned embedding model."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path("Bao_cao_danh_gia_model_embedding_finetune.docx")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()

    title = doc.add_heading("BÁO CÁO ĐÁNH GIÁ MODEL EMBEDDING FINE-TUNE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Semantic Product Search tiếng Việt - Phiên bản MVP (MiniLM)")
    doc.add_paragraph("Dự án: LLM Provider Benchmarking / Embedding Project")
    doc.add_paragraph("Ngày báo cáo: 28/05/2026")

    doc.add_heading("1. Tóm tắt điều hành", level=1)
    doc.add_paragraph(
        "Báo cáo đánh giá kết quả fine-tune embedding model đầu tiên cho semantic product search "
        "tiếng Việt. Model nền: sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2. "
        "Huấn luyện bằng MultipleNegativesRankingLoss trên cặp query-positive. "
        "Đánh giá gồm metric retrieval tự động và đánh giá thủ công A/B trên 30 query thực tế."
    )

    doc.add_heading("2. Phạm vi đánh giá", level=1)
    for item in [
        "Loại model: Embedding model (Sentence Transformers), không phải LLM sinh văn bản.",
        "Mục tiêu: map query người dùng sang searchable_text sản phẩm.",
        "Corpus: khoảng 2.000 sản phẩm (merged_products_vi_cleaned.csv).",
        "Dữ liệu: train ~6.224, valid ~778, test ~779 cặp query-positive.",
        "Model output: embedding_project/models/minilm_finetuned_final/",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Lưu ý: Training Loss/Perplexity của LLM (ví dụ Qwen2.5) không áp dụng trực tiếp cho "
        "pipeline embedding này. Nếu triển khai giai đoạn 2 (sinh mô tả bằng LLM), các metric đó "
        "sẽ được bổ sung trong báo cáo riêng."
    )

    doc.add_heading("3. Bộ metric sử dụng", level=1)
    doc.add_heading("3.1. Metric retrieval tự động (Top-K = 10)", level=2)
    add_table(
        doc,
        ["Metric", "Ý nghĩa"],
        [
            ["NDCG@10", "Đo chất lượng xếp hạng, ưu tiên relevant ở top."],
            ["Precision@10", "Tỷ lệ kết quả relevant trong top 10."],
            ["Recall@10 / Hit Rate@10", "Tỷ lệ query có ít nhất 1 relevant trong top 10."],
            ["MRR@10", "Thứ hạng trung bình của relevant đầu tiên."],
        ],
    )

    doc.add_heading("3.2. Metric đánh giá thủ công (Top-1)", level=2)
    for item in [
        "Hit@1: % query có kết quả #1 relevant (label = 2).",
        "Bad@1: % query có kết quả #1 irrelevant (label = 0).",
        "AvgLabel@1: điểm trung bình 0/1/2.",
        "Paired win-rate: % query fine-tuned có điểm cao hơn pretrained.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Quá trình huấn luyện (Training / Validation Loss)", level=1)
    add_table(
        doc,
        ["Epoch", "Training Loss", "Validation Loss"],
        [
            ["1", "0.077253", "0.045055"],
            ["2", "0.051476", "0.029501"],
        ],
    )
    doc.add_paragraph("Nhận xét quá trình train:")
    for item in [
        "Training loss giảm từ 0.0773 xuống 0.0515 (-33.3% so với epoch 1).",
        "Validation loss giảm từ 0.0451 xuống 0.0295 (-34.5% so với epoch 1).",
        "Không có dấu hiệu overfit rõ: validation loss giảm đồng thuận với training loss.",
        "Loss validation thấp hơn training loss ở cả 2 epoch — mô hình học ổn định trên tập valid.",
        "Đây là loss của contrastive training (MultipleNegativesRankingLoss), không phải perplexity LLM.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Kết quả đánh giá tự động (Retrieval)", level=1)
    add_table(
        doc,
        ["Metric", "Pretrained", "Fine-tuned"],
        [
            ["NDCG@10", "0.313", "0.778"],
            ["Precision@10", "0.056", "0.125"],
            ["Recall@10 (Hit Rate@10)", "0.434", "0.923"],
            ["MRR@10", "0.289", "0.736"],
        ],
    )

    doc.add_paragraph("Nhận xét:")
    for item in [
        "Fine-tuned cải thiện mạnh Recall@10/Hit Rate@10.",
        "NDCG@10 tăng từ 0.313 lên 0.778.",
        "Precision@10 vẫn thấp (12.5%), top-10 còn nhiễu.",
        "Cần đối chiếu thêm đánh giá thủ công để tránh overfit nhãn synthetic.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Kết quả đánh giá thủ công A/B (30 query)", level=1)
    add_table(
        doc,
        ["Metric", "Pretrained", "Fine-tuned"],
        [
            ["Hit@1", "50.0%", "73.3%"],
            ["Bad@1", "40.0%", "10.0%"],
            ["AvgLabel@1", "1.10", "1.63"],
            ["PartialOrBetter@1", "60.0%", "90.0%"],
            ["Paired wins", "1 query", "12 query"],
            ["Tie", "17 query", "17 query"],
        ],
    )

    doc.add_heading("7. So sánh với kỳ vọng", level=1)
    add_table(
        doc,
        ["Chỉ tiêu", "Mục tiêu", "Kết quả", "Đánh giá"],
        [
            ["NDCG@10", "> 0.85", "0.778", "Gần đạt"],
            ["Hit Rate@10", "> 90%", "92.3%", "Đạt (auto test)"],
            ["Hit@1 manual", "Cao", "73.3%", "Khả dụng MVP"],
            ["Precision@10", "Cao", "12.5%", "Chưa đạt"],
        ],
    )

    doc.add_heading("8. Hạn chế", level=1)
    for item in [
        "Nhãn train/query phần lớn sinh tự động.",
        "Đánh giá thủ công mới 30 query.",
        "Precision@10 thấp.",
        "Chưa có hard negative mining và hybrid rerank.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Hướng cải tiến", level=1)
    doc.add_heading("9.1. Ngắn hạn (embedding)", level=2)
    for item in [
        "Thu thập 200-500 query thật, chấm tay.",
        "Hard negatives cùng category/brand.",
        "Tối ưu searchable_text.",
        "Rerank 2 tầng: embedding top-50 + cross-encoder.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9.2. Trung hạn (mở rộng)", level=2)
    for item in [
        "PEFT (LoRA/QLoRA) cho LLM mô tả sản phẩm (Qwen2.5-Instruct).",
        "Contextual prompting: ghép trường title/category/brand/attributes có chọn lọc.",
        "A/B online theo CTR/conversion.",
        "Triển khai Qdrant + filter metadata.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10. Khuyến nghị", level=1)
    doc.add_paragraph(
        "Dùng fine-tuned làm model retrieval chính cho MVP, giữ pretrained làm fallback. "
        "Chưa production-final cho đến khi Precision@5/10 và đánh giá thủ công mở rộng đạt ngưỡng."
    )

    doc.add_heading("11. Phụ lục - Cấu hình huấn luyện", level=1)
    for item in [
        "Base: paraphrase-multilingual-MiniLM-L12-v2",
        "Loss: MultipleNegativesRankingLoss",
        "Epochs: 2-3",
        "Batch size: 8-16",
        "Learning rate: 2e-5",
        "Max seq length: 256-384",
        "Warmup: 0.1",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT.resolve()}")


if __name__ == "__main__":
    main()
